"""Generate professional PDF and DOCX research-style reports for the MRS project."""

from __future__ import annotations

import csv
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple

import joblib
import numpy as np


ROOT = Path(__file__).resolve().parent
DATASET_DIR = ROOT / "datasets"
ARTIFACT_PATH = ROOT / "artifacts" / "model_bundle.joblib"
PDF_PATH = ROOT / "Medicine_Recommendation_System_Research_Report.pdf"
DOCX_PATH = ROOT / "Medicine_Recommendation_System_Research_Report.docx"
PPTX_PATH = ROOT / "Medicine_Recommendation_System_Research_Presentation.pptx"


def ensure_package(import_name: str, pip_name: str) -> None:
    """Install package in current environment if missing."""
    try:
        __import__(import_name)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pip_name])


def count_csv_rows(file_path: Path) -> int:
    if not file_path.exists():
        return 0
    with file_path.open("r", encoding="utf-8-sig", newline="") as f:
        return max(sum(1 for _ in csv.reader(f)) - 1, 0)


def training_dataset_stats() -> Dict[str, int]:
    train_path = DATASET_DIR / "Training.csv"
    if not train_path.exists():
        return {"rows": 0, "symptoms": 0, "diseases": 0}

    with train_path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        fieldnames = reader.fieldnames or []
        symptom_cols = [c for c in fieldnames if c != "prognosis"]
        diseases = set()
        rows = 0
        for row in reader:
            rows += 1
            diseases.add((row.get("prognosis") or "").strip())

    return {"rows": rows, "symptoms": len(symptom_cols), "diseases": len(diseases)}


def normalize_text(value: str) -> str:
    return " ".join(str(value).replace("_", " ").strip().lower().split())


def read_column_values(csv_path: Path, column: str) -> List[str]:
    if not csv_path.exists():
        return []
    values: List[str] = []
    with csv_path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            raw = row.get(column)
            if raw is not None:
                values.append(str(raw))
    return values


def cross_file_label_mismatches() -> Dict[str, int]:
    training_labels = {
        normalize_text(v)
        for v in read_column_values(DATASET_DIR / "Training.csv", "prognosis")
        if str(v).strip()
    }

    files = [
        ("medications.csv", "Disease"),
        ("diets.csv", "Disease"),
        ("description.csv", "Disease"),
        ("precautions_df.csv", "Disease"),
    ]

    report: Dict[str, int] = {}
    for filename, col in files:
        other = {
            normalize_text(v)
            for v in read_column_values(DATASET_DIR / filename, col)
            if str(v).strip()
        }
        missing = training_labels - other
        if missing:
            report[filename] = len(missing)
    return report


def evaluate_from_bundle() -> Dict[str, object]:
    if not ARTIFACT_PATH.exists():
        raise FileNotFoundError(
            f"Model bundle not found at {ARTIFACT_PATH}. Run ml/train_model.py first."
        )

    bundle = joblib.load(ARTIFACT_PATH)
    model = bundle["model"]
    label_encoder = bundle["label_encoder"]
    feature_columns = bundle["feature_columns"]

    train_path = DATASET_DIR / "Training.csv"
    if not train_path.exists():
        raise FileNotFoundError(f"Training dataset not found: {train_path}")

    # Load with numpy to avoid importing pandas.
    with train_path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    X = np.zeros((len(rows), len(feature_columns)), dtype=np.int8)
    y_labels: List[str] = []
    for i, row in enumerate(rows):
        for j, col in enumerate(feature_columns):
            X[i, j] = int(float(row.get(col, 0) or 0))
        y_labels.append((row.get("prognosis") or "").strip())

    y_true = label_encoder.transform(y_labels)

    # Recreate the same split used in training.
    from sklearn.model_selection import train_test_split
    from sklearn.metrics import accuracy_score, f1_score, classification_report

    _, X_test, _, y_test = train_test_split(
        X,
        y_true,
        test_size=0.2,
        random_state=42,
        stratify=y_true,
    )

    y_pred = model.predict(X_test)
    y_prob = model.predict_proba(X_test)

    macro_f1 = float(f1_score(y_test, y_pred, average="macro"))
    accuracy = float(accuracy_score(y_test, y_pred))
    top3 = float(
        np.mean([
            y_true_i in np.argsort(prob_row)[-3:]
            for y_true_i, prob_row in zip(y_test, y_prob)
        ])
    )

    cls_report = classification_report(
        y_test,
        y_pred,
        target_names=label_encoder.classes_,
        output_dict=True,
        zero_division=0,
    )

    # Build a concrete example output from first test row.
    sample_vec = X_test[0:1]
    sample_prob = model.predict_proba(sample_vec)[0]
    sample_pred = int(model.predict(sample_vec)[0])
    sample_truth = int(y_test[0])
    top_idx = np.argsort(sample_prob)[::-1][:3]

    top_candidates = [
        {
            "disease": str(label_encoder.inverse_transform([idx])[0]),
            "probability": float(sample_prob[idx]),
        }
        for idx in top_idx
    ]

    active_symptoms = [
        feature_columns[i]
        for i, val in enumerate(sample_vec[0])
        if int(val) == 1
    ]

    return {
        "accuracy": accuracy,
        "macro_f1": macro_f1,
        "top3": top3,
        "weighted_precision": float(cls_report["weighted avg"]["precision"]),
        "weighted_recall": float(cls_report["weighted avg"]["recall"]),
        "n_classes": len(label_encoder.classes_),
        "test_size": int(len(y_test)),
        "sample": {
            "actual": str(label_encoder.inverse_transform([sample_truth])[0]),
            "predicted": str(label_encoder.inverse_transform([sample_pred])[0]),
            "top_candidates": top_candidates,
            "symptoms": active_symptoms[:10],
        },
    }


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def build_sections(dataset_stats: Dict[str, int], metrics: Dict[str, object], mismatches: Dict[str, int]) -> List[Tuple[str, str]]:
    mismatch_text = "None" if not mismatches else "; ".join(
        f"{k}: {v}" for k, v in sorted(mismatches.items())
    )

    sample = metrics["sample"]
    sample_candidates = "; ".join(
        f"{c['disease']} ({c['probability'] * 100:.2f}%)"
        for c in sample["top_candidates"]
    )

    sections: List[Tuple[str, str]] = [
        (
            "Abstract",
            clean_text(
                "This paper presents a Medicine Recommendation System (MRS) that predicts probable diseases from user symptoms and "
                "provides supportive recommendations such as medications, precautions, dietary guidance, and workout suggestions. "
                "A Random Forest classifier is trained on curated symptom-disease data and integrated into a Django web application. "
                f"On the held-out test set (n={metrics['test_size']}), the system achieved Accuracy={metrics['accuracy']:.4f}, "
                f"Macro F1={metrics['macro_f1']:.4f}, and Top-3 Accuracy={metrics['top3']:.4f}."
            ),
        ),
        (
            "1. Problem Definition",
            clean_text(
                "Early symptom-based disease triage is often delayed by limited expert availability and fragmented patient information. "
                "The project objective is to build a decision-support pipeline that converts symptom inputs into probable diseases and "
                "actionable guidance while clearly preserving physician oversight."
            ),
        ),
        (
            "2. Dataset Used",
            clean_text(
                "The system uses local datasets from datasets/: Training.csv, symptoms_df.csv, medications.csv, description.csv, "
                "precautions_df.csv, Symptom-severity.csv, diets.csv, and workout_df.csv. "
                f"Training.csv contains {dataset_stats['rows']} samples, {dataset_stats['symptoms']} symptom features, and "
                f"{dataset_stats['diseases']} unique diseases. Label consistency checks identified mismatches: {mismatch_text}."
            ),
        ),
        (
            "3. Model Trained",
            clean_text(
                "A RandomForestClassifier was trained with n_estimators=300, class_weight=balanced_subsample, random_state=42, and n_jobs=-1. "
                "Data split strategy: stratified train-test split with 80% training and 20% testing. "
                "The artifact includes model, label encoder, feature columns, and symptom severity map for reproducible inference."
            ),
        ),
        (
            "4. Accuracy and Performance",
            clean_text(
                f"Evaluation on held-out test data produced Accuracy={metrics['accuracy']:.4f}, Macro F1={metrics['macro_f1']:.4f}, "
                f"Top-3 Accuracy={metrics['top3']:.4f}, Weighted Precision={metrics['weighted_precision']:.4f}, "
                f"Weighted Recall={metrics['weighted_recall']:.4f}, across {metrics['n_classes']} disease classes."
            ),
        ),
        (
            "5. Output Visibility (Example)",
            clean_text(
                f"Sample inference from the evaluated test set: Actual disease={sample['actual']}; Predicted disease={sample['predicted']}. "
                f"Top-3 candidates: {sample_candidates}. Input symptoms (subset): {', '.join(sample['symptoms']) if sample['symptoms'] else 'N/A'}."
            ),
        ),
        (
            "6. System Output and Presentation",
            clean_text(
                "The Django UI collects age, gender, and multi-symptom inputs; then returns disease confidence, top candidates, "
                "description, medication suggestions, precautions, diet, and workout advice. Responses are rendered through a "
                "structured result page and JSON endpoints for asynchronous requests."
            ),
        ),
        (
            "7. Discussion",
            clean_text(
                "Perfect test metrics indicate strong separability in the present dataset. However, practical deployment requires external "
                "validation, calibration checks, and richer clinical covariates to reduce overfitting risk and improve robustness."
            ),
        ),
        (
            "8. Conclusion",
            clean_text(
                "The MRS pipeline demonstrates a complete end-to-end framework: data ingestion, model training, reproducible artifacts, "
                "web-based inference, and recommendation display. It should be used strictly as educational decision support, not as a "
                "replacement for professional medical diagnosis."
            ),
        ),
    ]

    return sections


def generate_docx(title: str, authors: str, sections: List[Tuple[str, str]], metrics: Dict[str, object], dataset_stats: Dict[str, int]) -> None:
    ensure_package("docx", "python-docx")
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(authors)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(datetime.now().strftime("%B %d, %Y"))

    doc.add_paragraph("")

    for heading, body in sections:
        h = doc.add_paragraph()
        hr = h.add_run(heading)
        hr.bold = True
        hr.font.size = Pt(12)

        bp = doc.add_paragraph(body)
        bp.paragraph_format.space_after = Pt(8)

    # Add compact metric table for professional format.
    doc.add_paragraph("")
    h = doc.add_paragraph()
    hr = h.add_run("Appendix A: Core Metrics")
    hr.bold = True
    hr.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"

    rows = [
        ("Training Samples", str(dataset_stats["rows"])),
        ("Symptom Features", str(dataset_stats["symptoms"])),
        ("Disease Classes", str(metrics["n_classes"])),
        ("Test Samples", str(metrics["test_size"])),
        ("Accuracy", f"{metrics['accuracy']:.4f}"),
        ("Macro F1", f"{metrics['macro_f1']:.4f}"),
        ("Top-3 Accuracy", f"{metrics['top3']:.4f}"),
        ("Weighted Precision", f"{metrics['weighted_precision']:.4f}"),
        ("Weighted Recall", f"{metrics['weighted_recall']:.4f}"),
    ]

    for key, val in rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = val

    doc.add_paragraph("")
    d = doc.add_paragraph(
        "Disclaimer: This report is for educational and research presentation purposes only. "
        "The system is a decision-support tool and not a substitute for licensed clinical judgment."
    )
    d.runs[0].italic = True

    doc.save(DOCX_PATH)


def generate_pdf(title: str, authors: str, sections: List[Tuple[str, str]], metrics: Dict[str, object], dataset_stats: Dict[str, int]) -> None:
    ensure_package("fpdf", "fpdf2")
    from fpdf import FPDF

    pdf = FPDF(format="letter")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def heading(text: str, size: int = 12) -> None:
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Times", "B", size)
        pdf.multi_cell(0, 7, text)

    def body(text: str) -> None:
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Times", "", 11)
        pdf.multi_cell(0, 6, text)
        pdf.ln(1)

    pdf.set_font("Times", "B", 18)
    pdf.multi_cell(0, 10, title, align="C")
    pdf.set_font("Times", "", 11)
    pdf.cell(0, 6, authors, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.cell(0, 6, datetime.now().strftime("%B %d, %Y"), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(3)

    for section_title, section_body in sections:
        heading(section_title, size=12)
        body(section_body)

    heading("Appendix A: Core Metrics", size=12)

    metric_lines = [
        f"Training Samples: {dataset_stats['rows']}",
        f"Symptom Features: {dataset_stats['symptoms']}",
        f"Disease Classes: {metrics['n_classes']}",
        f"Test Samples: {metrics['test_size']}",
        f"Accuracy: {metrics['accuracy']:.4f}",
        f"Macro F1: {metrics['macro_f1']:.4f}",
        f"Top-3 Accuracy: {metrics['top3']:.4f}",
        f"Weighted Precision: {metrics['weighted_precision']:.4f}",
        f"Weighted Recall: {metrics['weighted_recall']:.4f}",
    ]
    body("\n".join(metric_lines))

    pdf.set_font("Times", "I", 10)
    pdf.multi_cell(
        0,
        5,
        "Disclaimer: This report is for educational and research presentation purposes only. "
        "The system is a decision-support tool and not a substitute for licensed clinical judgment.",
    )

    pdf.output(str(PDF_PATH))


def generate_pptx(title: str, authors: str, metrics: Dict[str, object], dataset_stats: Dict[str, int], mismatches: Dict[str, int]) -> None:
    ensure_package("pptx", "python-pptx")

    import importlib

    Presentation = importlib.import_module("pptx").Presentation
    ChartData = importlib.import_module("pptx.chart.data").ChartData
    RGBColor = importlib.import_module("pptx.dml.color").RGBColor
    XL_CHART_TYPE = importlib.import_module("pptx.enum.chart").XL_CHART_TYPE
    MSO_AUTO_SHAPE_TYPE = importlib.import_module("pptx.enum.shapes").MSO_AUTO_SHAPE_TYPE
    Inches = importlib.import_module("pptx.util").Inches
    Pt = importlib.import_module("pptx.util").Pt

    prs = Presentation()
    prs.slide_width = Inches(13.333)
    prs.slide_height = Inches(7.5)

    bg_color = RGBColor(247, 250, 252)
    panel_color = RGBColor(255, 255, 255)
    title_color = RGBColor(15, 41, 77)
    accent_color = RGBColor(0, 116, 217)
    accent_soft = RGBColor(224, 238, 255)
    text_color = RGBColor(40, 52, 68)

    def add_background(slide) -> None:
        bg = slide.shapes.add_shape(MSO_AUTO_SHAPE_TYPE.RECTANGLE, Inches(0), Inches(0), prs.slide_width, prs.slide_height)
        bg.fill.solid()
        bg.fill.fore_color.rgb = bg_color
        bg.line.fill.background()
        bar = slide.shapes.add_shape(MSO_AUTO_SHAPE_TYPE.RECTANGLE, Inches(0), Inches(0), prs.slide_width, Inches(0.35))
        bar.fill.solid()
        bar.fill.fore_color.rgb = accent_color
        bar.line.fill.background()
        orb = slide.shapes.add_shape(MSO_AUTO_SHAPE_TYPE.OVAL, Inches(11.6), Inches(-0.8), Inches(3), Inches(3))
        orb.fill.solid()
        orb.fill.fore_color.rgb = accent_soft
        orb.line.fill.background()

    def add_title_block(slide, heading: str, subtitle: str | None = None) -> None:
        title_box = slide.shapes.add_textbox(Inches(0.8), Inches(0.55), Inches(11.8), Inches(1.1))
        tf = title_box.text_frame
        tf.clear()
        p = tf.paragraphs[0]
        p.text = heading
        p.font.name = "Segoe UI Semibold"
        p.font.size = Pt(32)
        p.font.color.rgb = title_color
        if subtitle:
            p2 = tf.add_paragraph()
            p2.text = subtitle
            p2.font.name = "Segoe UI"
            p2.font.size = Pt(16)
            p2.font.color.rgb = RGBColor(85, 98, 112)

    def add_panel(slide, left: float, top: float, width: float, height: float):
        panel = slide.shapes.add_shape(
            MSO_AUTO_SHAPE_TYPE.ROUNDED_RECTANGLE,
            Inches(left),
            Inches(top),
            Inches(width),
            Inches(height),
        )
        panel.fill.solid()
        panel.fill.fore_color.rgb = panel_color
        panel.line.color.rgb = accent_soft
        return panel

    def add_bullets(slide, items: List[str], left: float, top: float, width: float, height: float, font_size: int = 21) -> None:
        box = slide.shapes.add_textbox(Inches(left), Inches(top), Inches(width), Inches(height))
        tf = box.text_frame
        tf.clear()
        for idx, item in enumerate(items):
            para = tf.paragraphs[0] if idx == 0 else tf.add_paragraph()
            para.text = item
            para.level = 0
            para.font.name = "Segoe UI"
            para.font.size = Pt(font_size)
            para.font.color.rgb = text_color
            para.space_after = Pt(12)

    sample = metrics["sample"]
    sample_candidates = [
        f"{c['disease']} ({c['probability'] * 100:.1f}%)"
        for c in sample["top_candidates"]
    ]

    dataset_files = [
        "Training.csv",
        "symptoms_df.csv",
        "medications.csv",
        "description.csv",
        "precautions_df.csv",
        "Symptom-severity.csv",
        "diets.csv",
        "workout_df.csv",
    ]

    mismatch_text = "None" if not mismatches else "; ".join(f"{k}: {v}" for k, v in sorted(mismatches.items()))

    # Slide 1: Cover
    slide = prs.slides.add_slide(prs.slide_layouts[6])
    add_background(slide)
    add_title_block(slide, "Medicine Recommendation System", "Machine Learning-Powered Clinical Decision Support")
    add_panel(slide, 0.8, 2.0, 7.7, 2.7)
    cover_lines = [
        "Research Presentation",
        authors,
        datetime.now().strftime("%B %d, %Y"),
    ]
    add_bullets(slide, cover_lines, 1.1, 2.35, 7.0, 2.0, font_size=24)

    # Slide 2: Abstract
    slide = prs.slides.add_slide(prs.slide_layouts[6])
    add_background(slide)
    add_title_block(slide, "Abstract")
    add_panel(slide, 0.8, 1.8, 11.8, 4.9)
    abstract_points = [
        "Predicts probable diseases from symptom inputs and provides medication, diet, precautions, and workout guidance.",
        "Combines Random Forest inference with a structured medical knowledge base in a Django web application.",
        f"Performance on held-out test set: Accuracy {metrics['accuracy']:.4f}, Macro F1 {metrics['macro_f1']:.4f}, Top-3 {metrics['top3']:.4f}.",
        "Designed as educational decision support with clear physician-oversight disclaimer.",
    ]
    add_bullets(slide, abstract_points, 1.1, 2.15, 11.1, 4.2, font_size=20)

    # Slide 3: Datasets
    slide = prs.slides.add_slide(prs.slide_layouts[6])
    add_background(slide)
    add_title_block(slide, "Datasets Used")
    add_panel(slide, 0.8, 1.7, 6.0, 5.0)
    add_panel(slide, 7.0, 1.7, 5.6, 5.0)

    file_lines = [
        f"{name}: {count_csv_rows(DATASET_DIR / name)} rows"
        for name in dataset_files
    ]
    add_bullets(slide, file_lines, 1.1, 2.0, 5.5, 4.6, font_size=16)

    rhs_points = [
        f"Training samples: {dataset_stats['rows']}",
        f"Symptom features: {dataset_stats['symptoms']}",
        f"Disease classes: {dataset_stats['diseases']}",
        f"Label mismatch check: {mismatch_text}",
    ]
    add_bullets(slide, rhs_points, 7.3, 2.0, 5.0, 4.6, font_size=18)

    # Slide 4: Model
    slide = prs.slides.add_slide(prs.slide_layouts[6])
    add_background(slide)
    add_title_block(slide, "Model Used")
    add_panel(slide, 0.8, 1.8, 11.8, 4.9)
    model_points = [
        "Algorithm: RandomForestClassifier",
        "Configuration: n_estimators=300, class_weight=balanced_subsample, random_state=42, n_jobs=-1",
        "Data strategy: stratified 80/20 train-test split",
        "Saved artifacts: model, label encoder, feature columns, severity map",
        "Serving layer: Django endpoint + CSV-driven recommendation enrichment",
    ]
    add_bullets(slide, model_points, 1.1, 2.15, 11.1, 4.2, font_size=19)

    # Slide 5: Accuracy
    slide = prs.slides.add_slide(prs.slide_layouts[6])
    add_background(slide)
    add_title_block(slide, "Accuracy and Performance")
    add_panel(slide, 0.8, 1.8, 5.1, 4.9)
    add_panel(slide, 6.1, 1.8, 6.5, 4.9)

    metric_lines = [
        f"Accuracy: {metrics['accuracy']:.4f}",
        f"Macro F1: {metrics['macro_f1']:.4f}",
        f"Top-3 Accuracy: {metrics['top3']:.4f}",
        f"Weighted Precision: {metrics['weighted_precision']:.4f}",
        f"Weighted Recall: {metrics['weighted_recall']:.4f}",
        f"Test samples: {metrics['test_size']}",
    ]
    add_bullets(slide, metric_lines, 1.1, 2.1, 4.6, 4.2, font_size=17)

    chart_data = ChartData()
    chart_data.categories = ["Accuracy", "Macro F1", "Top-3", "W. Precision", "W. Recall"]
    chart_data.add_series(
        "Score (%)",
        [
            metrics["accuracy"] * 100,
            metrics["macro_f1"] * 100,
            metrics["top3"] * 100,
            metrics["weighted_precision"] * 100,
            metrics["weighted_recall"] * 100,
        ],
    )
    chart = slide.shapes.add_chart(
        XL_CHART_TYPE.COLUMN_CLUSTERED,
        Inches(6.45),
        Inches(2.2),
        Inches(5.8),
        Inches(4.2),
        chart_data,
    ).chart
    chart.has_legend = False
    chart.value_axis.maximum_scale = 100
    chart.value_axis.minimum_scale = 0
    chart.value_axis.has_major_gridlines = True

    # Slide 6: Output
    slide = prs.slides.add_slide(prs.slide_layouts[6])
    add_background(slide)
    add_title_block(slide, "Output Visible")
    add_panel(slide, 0.8, 1.8, 11.8, 4.9)

    output_points = [
        f"Input symptoms: {', '.join(sample['symptoms']) if sample['symptoms'] else 'N/A'}",
        f"Predicted disease: {sample['predicted']}",
        f"Actual disease in test sample: {sample['actual']}",
        f"Top-3 candidates: {'; '.join(sample_candidates)}",
        "UI output includes confidence, disease description, medication list, precautions, diet, and workout guidance.",
    ]
    add_bullets(slide, output_points, 1.1, 2.15, 11.1, 4.2, font_size=18)

    # Slide 7: System Flow
    slide = prs.slides.add_slide(prs.slide_layouts[6])
    add_background(slide)
    add_title_block(slide, "System Workflow")
    add_panel(slide, 0.8, 1.8, 11.8, 4.9)
    flow_points = [
        "Patient enters age, gender, and symptoms on web form.",
        "Backend converts symptoms to binary feature vector.",
        "Random Forest predicts disease and confidence scores.",
        "Knowledge base maps disease to medications, precautions, diet, and workouts.",
        "Result page / JSON response presents complete recommendation package.",
    ]
    add_bullets(slide, flow_points, 1.1, 2.15, 11.1, 4.2, font_size=19)

    # Slide 8: Conclusion
    slide = prs.slides.add_slide(prs.slide_layouts[6])
    add_background(slide)
    add_title_block(slide, "Conclusion")
    add_panel(slide, 0.8, 1.8, 11.8, 4.9)
    conclusion_points = [
        "MRS demonstrates an end-to-end ML + web recommendation pipeline with reproducible artifacts.",
        "The system achieves strong benchmark performance on the current dataset.",
        "Future work: external validation, calibration studies, and expanded clinical features.",
        "Use case: educational and preliminary triage support, not a replacement for clinical diagnosis.",
    ]
    add_bullets(slide, conclusion_points, 1.1, 2.15, 11.1, 4.2, font_size=20)

    prs.save(PPTX_PATH)


def main() -> None:
    title = "Medicine Recommendation System: A Machine Learning Decision-Support Framework"
    authors = "Project Report | Django + Random Forest"

    dataset_stats = training_dataset_stats()
    mismatches = cross_file_label_mismatches()
    metrics = evaluate_from_bundle()
    sections = build_sections(dataset_stats, metrics, mismatches)

    generate_pdf(title, authors, sections, metrics, dataset_stats)
    generate_docx(title, authors, sections, metrics, dataset_stats)
    generate_pptx(title, authors, metrics, dataset_stats, mismatches)

    print("Generated documents:")
    print(f"- PDF : {PDF_PATH}")
    print(f"- DOCX: {DOCX_PATH}")
    print(f"- PPTX: {PPTX_PATH}")


if __name__ == "__main__":
    main()
